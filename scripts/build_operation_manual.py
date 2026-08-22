from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output"
ASSETS = OUT / "manual-assets"
DOCX = OUT / "数智课堂教育智能体操作手册_V1.0.docx"
OUT.mkdir(exist_ok=True)
ASSETS.mkdir(exist_ok=True)

NAVY = "123B57"
CYAN = "0D8CA3"
TEAL = "19A7A0"
LIGHT = "EAF6F8"
PALE = "F5F8FA"
GRAY = "64748B"
INK = "1F2937"
RED = "B42318"
GOLD = "B7791F"
WHITE = "FFFFFF"


def font_path():
    choices = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
    ]
    return next((p for p in choices if Path(p).exists()), None)


FONT_FILE = font_path()


def pil_font(size, bold=False):
    if FONT_FILE:
        return ImageFont.truetype(FONT_FILE, size=size, index=1 if bold else 0)
    return ImageFont.load_default()


def rounded_box(draw, xy, text, fill=WHITE, outline="#A8C7CF", size=30, bold=False, radius=18):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    f = pil_font(size, bold)
    box = draw.multiline_textbbox((0, 0), text, font=f, spacing=6, align="center")
    w, h = box[2] - box[0], box[3] - box[1]
    draw.multiline_text(((x1+x2-w)/2, (y1+y2-h)/2), text, font=f, fill="#17324A", spacing=6, align="center")


def arrow(draw, p1, p2, color="#5B7285", width=4, dashed=False):
    x1, y1 = p1; x2, y2 = p2
    if dashed:
        steps = 18
        for i in range(0, steps, 2):
            a, b = i/steps, min((i+1)/steps, 1)
            draw.line((x1+(x2-x1)*a, y1+(y2-y1)*a, x1+(x2-x1)*b, y1+(y2-y1)*b), fill=color, width=width)
    else:
        draw.line((x1, y1, x2, y2), fill=color, width=width)
    import math
    ang = math.atan2(y2-y1, x2-x1)
    s = 16
    pts = [(x2, y2), (x2-s*math.cos(ang-.55), y2-s*math.sin(ang-.55)), (x2-s*math.cos(ang+.55), y2-s*math.sin(ang+.55))]
    draw.polygon(pts, fill=color)


def make_agent_flow(path):
    im = Image.new("RGB", (1800, 1120), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 45), "数智课堂多Agent协同与长期学习记忆流程", font=pil_font(44, True), fill="#123B57")
    boxes = {
        "u": (620, 130, 1180, 225), "o": (650, 280, 1150, 385),
        "p": (110, 470, 430, 570), "e": (530, 470, 850, 570), "t": (950, 470, 1270, 570), "k": (1370, 470, 1690, 570),
        "q": (1140, 675, 1470, 775), "a": (720, 675, 1040, 775), "f": (300, 675, 620, 775),
        "m1": (260, 900, 670, 1015), "m2": (790, 900, 1240, 1015), "m3": (1360, 900, 1700, 1015),
    }
    rounded_box(d, boxes["u"], "教师或学生提出教学需求", fill="#F5F8FA", size=31, bold=True)
    rounded_box(d, boxes["o"], "小智总调度Agent\n识别意图并分配任务", fill="#DDF6F6", outline="#19A7A0", size=30, bold=True)
    rounded_box(d, boxes["p"], "理解规划Agent\n匹配模型、生成步骤", fill="#EEF2FF", size=27, bold=True)
    rounded_box(d, boxes["e"], "演示执行Agent\n调用受控教学工具", fill="#ECFDF5", size=27, bold=True)
    rounded_box(d, boxes["t"], "3D模型与交互工具\n加载·旋转·缩放·拆解", fill="#F8FAFC", size=25)
    rounded_box(d, boxes["k"], "知识讲解Agent\n生成并播报讲解", fill="#FFF7E6", size=27, bold=True)
    rounded_box(d, boxes["q"], "活泼追问Agent\n生成可选课堂追问", fill="#FDF2F8", size=27, bold=True)
    rounded_box(d, boxes["a"], "学生作答\n点击·语音·手势", fill="#F8FAFC", size=27)
    rounded_box(d, boxes["f"], "判题、解释与\n鼓励性反馈", fill="#F0FDFA", size=27)
    rounded_box(d, boxes["m1"], "保存本轮会话记录\n用户输入·回复·讲解·答题事件", fill="#F8FAFC", size=24)
    rounded_box(d, boxes["m2"], "异步整理\n会话摘要＋五类长期学习记忆", fill="#EAF6F8", outline="#0D8CA3", size=26, bold=True)
    rounded_box(d, boxes["m3"], "只读学习背景\n辅助总调度调整讲解与复习建议", fill="#F0FDFA", outline="#19A7A0", size=24)
    arrow(d, (900,225), (900,280)); arrow(d,(650,350),(270,470)); arrow(d,(430,520),(530,520)); arrow(d,(850,520),(950,520)); arrow(d,(1270,520),(1370,520))
    arrow(d,(1530,570),(1305,675)); arrow(d,(1140,725),(1040,725)); arrow(d,(720,725),(620,725))
    arrow(d,(820,385),(1305,675), color="#A35B86")
    arrow(d,(460,775),(465,900), dashed=True); arrow(d,(880,775),(520,900), dashed=True); arrow(d,(670,958),(790,958), dashed=True); arrow(d,(1240,958),(1360,958), dashed=True)
    arrow(d,(1530,900),(1140,350), color="#19A7A0", dashed=True)
    d.text((80, 1060), "实线：课堂任务执行路径    虚线：学习记忆保存、整理与只读引用路径", font=pil_font(24), fill="#64748B")
    im.save(path, quality=95)


def make_sequence(path):
    im = Image.new("RGB", (1800, 940), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 40), "典型课堂任务时序：讲解地球内部结构", font=pil_font(44, True), fill="#123B57")
    lanes = [(90,"教师"),(370,"小智总调度"),(680,"理解规划"),(990,"演示执行"),(1300,"知识讲解"),(1590,"活泼追问")]
    for x, name in lanes:
        d.text((x-45,130), name, font=pil_font(25, True), fill="#17324A")
        d.line((x,180,x,850), fill="#CBD5E1", width=3)
    events = [
        (220,90,370,"“讲解地球内部结构”"),(300,370,680,"识别教学演示任务"),(390,680,990,"模型＋结构化步骤"),
        (480,990,990,"加载、缩放、旋转、拆解"),(570,990,1300,"当前模型与教学目标"),(660,1300,90,"讲解内容＋语音播报"),
        (750,1300,1590,"可选：生成课堂追问"),(830,1590,90,"题目、判题与解释反馈"),
    ]
    for y,x1,x2,label in events:
        arrow(d,(x1,y),(x2,y), color="#0D8CA3")
        bbox=d.textbbox((0,0),label,font=pil_font(21)); w=bbox[2]
        d.rectangle(((x1+x2-w)/2-8,y-33,(x1+x2+w)/2+8,y-5),fill="white")
        d.text(((x1+x2-w)/2,y-31),label,font=pil_font(21),fill="#334155")
    im.save(path, quality=95)


def make_workspace(path):
    im = Image.new("RGB", (1800, 950), "white")
    d = ImageDraw.Draw(im)
    d.text((70, 40), "课堂工作台功能分区示意", font=pil_font(44, True), fill="#123B57")
    d.rounded_rectangle((70,130,1730,860),radius=28,fill="#F7FAFC",outline="#9FC6CF",width=4)
    d.rounded_rectangle((100,165,1700,260),radius=18,fill="#123B57")
    d.text((140,190),"数智课堂",font=pil_font(34,True),fill="white")
    d.text((1110,190),"3D建模生成    导入模型    个人中心",font=pil_font(25),fill="#DDF6F6")
    rounded_box(d,(105,300,430,800),"学科资源库\n\n我的模型\n建模资源\n化学模型\n生物模型\n地理模型",fill="#EAF6F8",size=28,bold=True)
    rounded_box(d,(470,300,1320,800),"3D模型展示区\n\n模型加载与观察\n手势/鼠标交互\n知识讲解\n全屏与答题模式",fill="white",size=30,bold=True)
    rounded_box(d,(1360,300,1695,540),"小智助教\n语音输入与状态",fill="#F0FDFA",size=27,bold=True)
    rounded_box(d,(1360,575,1695,800),"多智能体协作台\n状态·步骤·工具结果",fill="#FFF7E6",size=25,bold=True)
    im.save(path, quality=95)


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn(f'w:{m}'))
        if node is None: node=OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); tblHeader = OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)


def set_run_font(run, name="Noto Sans CJK SC", size=None, bold=None, color=None):
    run.font.name=name; run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'),name)
    if size: run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if color: run.font.color.rgb=RGBColor.from_string(color)


def page_number(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run=paragraph.add_run(); fldChar1=OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin')
    instr=OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text=' PAGE '
    fldChar2=OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end')
    run._r.extend([fldChar1,instr,fldChar2]); set_run_font(run,size=9,color=GRAY)


def add_toc(doc):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    run=p.add_run(); begin=OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'),'begin')
    instr=OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text='TOC \\o "1-3" \\h \\z \\u'
    separate=OxmlElement('w:fldChar'); separate.set(qn('w:fldCharType'),'separate')
    text=OxmlElement('w:t'); text.text='打开文档后在目录上右键选择“更新域”即可刷新页码。'
    end=OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'),'end')
    run._r.extend([begin,instr,separate,text,end]); set_run_font(run,size=10,color=GRAY)


def add_heading(doc, text, level=1):
    doc._manual_num_id = None
    p=doc.add_paragraph(text, style=f'Heading {level}')
    return p


def add_body(doc, text, bold_lead=None):
    doc._manual_num_id = None
    p=doc.add_paragraph(style='Normal')
    if bold_lead and text.startswith(bold_lead):
        r=p.add_run(bold_lead); set_run_font(r,bold=True)
        r=p.add_run(text[len(bold_lead):]); set_run_font(r)
    else:
        r=p.add_run(text); set_run_font(r)
    return p


def _new_numbering_id(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn('w:abstractNumId'))) for x in numbering.findall(qn('w:abstractNum'))]
    num_ids = [int(x.get(qn('w:numId'))) for x in numbering.findall(qn('w:num'))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1
    abstract = OxmlElement('w:abstractNum'); abstract.set(qn('w:abstractNumId'), str(abstract_id))
    multi = OxmlElement('w:multiLevelType'); multi.set(qn('w:val'), 'singleLevel'); abstract.append(multi)
    lvl = OxmlElement('w:lvl'); lvl.set(qn('w:ilvl'), '0')
    start = OxmlElement('w:start'); start.set(qn('w:val'), '1'); lvl.append(start)
    num_fmt = OxmlElement('w:numFmt'); num_fmt.set(qn('w:val'), 'decimal'); lvl.append(num_fmt)
    lvl_text = OxmlElement('w:lvlText'); lvl_text.set(qn('w:val'), '%1.'); lvl.append(lvl_text)
    jc = OxmlElement('w:lvlJc'); jc.set(qn('w:val'), 'left'); lvl.append(jc)
    ppr = OxmlElement('w:pPr'); tabs = OxmlElement('w:tabs'); tab = OxmlElement('w:tab'); tab.set(qn('w:val'),'num'); tab.set(qn('w:pos'),'560'); tabs.append(tab); ppr.append(tabs)
    ind = OxmlElement('w:ind'); ind.set(qn('w:left'),'560'); ind.set(qn('w:hanging'),'280'); ppr.append(ind); lvl.append(ppr)
    abstract.append(lvl); numbering.append(abstract)
    num = OxmlElement('w:num'); num.set(qn('w:numId'), str(num_id)); aid=OxmlElement('w:abstractNumId'); aid.set(qn('w:val'),str(abstract_id)); num.append(aid); numbering.append(num)
    return num_id


def bullet(doc, text, numbered=False):
    if numbered:
        if getattr(doc, '_manual_num_id', None) is None:
            doc._manual_num_id = _new_numbering_id(doc)
        p=doc.add_paragraph(style='Normal')
        pPr=p._p.get_or_add_pPr(); numPr=OxmlElement('w:numPr'); ilvl=OxmlElement('w:ilvl'); ilvl.set(qn('w:val'),'0'); numId=OxmlElement('w:numId'); numId.set(qn('w:val'),str(doc._manual_num_id)); numPr.extend([ilvl,numId]); pPr.append(numPr)
    else:
        doc._manual_num_id = None
        p=doc.add_paragraph(style='List Bullet')
    r=p.add_run(text); set_run_font(r)
    return p


def callout(doc, title, text, kind='info'):
    doc._manual_num_id = None
    colors={'info':(LIGHT,CYAN),'warn':('FFF7E6',GOLD),'risk':('FFF1F2',RED),'success':('ECFDF5',TEAL)}
    fill,accent=colors[kind]
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; t.columns[0].width=Cm(16.2)
    cell=t.cell(0,0); set_cell_shading(cell,fill); set_cell_margins(cell,140,180,140,180)
    p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(3)
    r=p.add_run(title); set_run_font(r,bold=True,color=accent,size=10.5)
    p2=cell.add_paragraph(); p2.paragraph_format.space_after=Pt(0); r=p2.add_run(text); set_run_font(r,size=10)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)


def add_table(doc, headers, rows, widths=None):
    doc._manual_num_id = None
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'; t.autofit=False
    if widths is None: widths=[16.2/len(headers)]*len(headers)
    for i,(h,w) in enumerate(zip(headers,widths)):
        c=t.rows[0].cells[i]; c.width=Cm(w); set_cell_shading(c,NAVY); set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(h); set_run_font(r,bold=True,color=WHITE,size=9.5)
    set_repeat_table_header(t.rows[0])
    for row in rows:
        cells=t.add_row().cells
        for i,(val,w) in enumerate(zip(row,widths)):
            cells[i].width=Cm(w); set_cell_margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=cells[i].paragraphs[0]; r=p.add_run(str(val)); set_run_font(r,size=9.2)
            if i==0: set_cell_shading(cells[i],PALE); r.bold=True
    return t


def add_figure(doc, path, caption, width=16.2):
    doc._manual_num_id = None
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next=True
    p.add_run().add_picture(str(path),width=Cm(width))
    c=doc.add_paragraph(caption,style='Caption'); c.alignment=WD_ALIGN_PARAGRAPH.CENTER


def page_break(doc):
    doc._manual_num_id = None
    doc.add_page_break()


def build():
    agent_flow=ASSETS/'agent-memory-flow.png'; seq=ASSETS/'earth-sequence.png'; workspace=ASSETS/'workspace-map.png'
    make_agent_flow(agent_flow); make_sequence(seq); make_workspace(workspace)
    doc=Document(); sec=doc.sections[0]
    sec.page_width=Cm(21); sec.page_height=Cm(29.7); sec.top_margin=Cm(2.2); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.4); sec.right_margin=Cm(2.4); sec.header_distance=Cm(1.0); sec.footer_distance=Cm(1.0)
    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Noto Sans CJK SC'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC'); normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string(INK)
    normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.35
    for name,size,color,before,after in [('Heading 1',17,NAVY,16,8),('Heading 2',13,CYAN,12,6),('Heading 3',11.5,NAVY,9,4)]:
        s=styles[name]; s.font.name='Noto Sans CJK SC'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC'); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    styles['Caption'].font.name='Noto Sans CJK SC'; styles['Caption']._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC'); styles['Caption'].font.size=Pt(9); styles['Caption'].font.color.rgb=RGBColor.from_string(GRAY)
    for lname in ['List Bullet','List Number']:
        styles[lname].font.name='Noto Sans CJK SC'; styles[lname]._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC'); styles[lname].font.size=Pt(10.5); styles[lname].paragraph_format.space_after=Pt(4)
    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=header.add_run('数智课堂教育智能体操作手册  ·  V1.0'); set_run_font(r,size=9,color=GRAY)
    page_number(sec.footer.paragraphs[0])

    # Cover
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(48); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    logo=ROOT/'public/brand/smart-cube-tech/mark.svg'
    # SVG embedding is not reliable in LibreOffice; use text mark.
    r=p.add_run('◇'); set_run_font(r,size=34,color=TEAL,bold=True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(16)
    r=p.add_run('数智课堂'); set_run_font(r,size=18,color=CYAN,bold=True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(30)
    r=p.add_run('教育智能体操作手册'); set_run_font(r,size=30,color=NAVY,bold=True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('总调度Agent · 专业Agent · 长期学习记忆'); set_run_font(r,size=13,color=GRAY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(40); r=p.add_run('V1.0'); set_run_font(r,size=20,color=TEAL,bold=True)
    callout(doc,'文档用途','本手册用于数智课堂教育智能体的申报、演示和教师使用培训，说明系统访问方式、智能体协作机制、长期学习记忆及主要课堂操作。')
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(35); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('发布日期：2026年7月    |    文档状态：申报版'); set_run_font(r,size=10,color=GRAY)

    page_break(doc)
    add_heading(doc,'文档信息',1)
    add_table(doc,['项目','内容'],[
        ('文档名称','数智课堂教育智能体操作手册'),('版本','V1.0'),('适用对象','教师、项目评审人员、系统管理员'),('正式访问地址','【请填写正式访问地址】'),('国内部署平台','【请填写国内部署平台名称】'),('推荐浏览器','最新版 Chrome 或 Microsoft Edge'),('主要能力','多Agent协作、3D教学、语音与手势、课堂答题、长期学习记忆')
    ],[4.0,12.2])
    add_heading(doc,'阅读说明',2)
    add_body(doc,'本手册中的“Agent”指承担特定教学职责的软件智能体；“工具”指由Agent调用的受控3D模型或交互能力。长期学习记忆为普通用户账号提供，管理员后台不进入学习记忆流程。')
    callout(doc,'重要边界','长期学习记忆用于辅助总调度Agent调整讲解难度和复习建议，不直接改变工具权限，不自动决定成绩，也不等同于正式学情评价。','warn')
    add_heading(doc,'目录',1); add_toc(doc)

    page_break(doc)
    add_heading(doc,'1 产品概述',1)
    add_heading(doc,'1.1 建设背景',2)
    add_body(doc,'数智课堂是一套面向课堂教学的多模态3D教育智能体系统。系统将自然语言理解、多Agent协作、三维模型演示、语音交互、空间手势和课堂答题结合起来，帮助教师把抽象知识转化为可观察、可操作、可追问的教学过程。')
    add_heading(doc,'1.2 适用对象',2)
    bullet(doc,'教师：选择教学模型、发起智能讲解、使用语音和手势控制课堂演示。')
    bullet(doc,'学生：观察3D模型，通过点击、语音或手势参与课堂答题。')
    bullet(doc,'评审人员：查看智能体分工、工具调用、长期记忆与教学闭环。')
    bullet(doc,'管理员：管理用户角色和账号启用状态。')
    add_heading(doc,'1.3 核心能力',2)
    add_table(doc,['能力','说明','教学价值'],[
        ('多Agent协作','总调度、规划、执行、讲解、追问分工协作','把复杂教学任务拆分为可观察步骤'),
        ('3D教学演示','模型加载、旋转、缩放、拆解、复位','强化空间结构认知'),
        ('多模态交互','鼠标、触控、语音、手势并行可用','降低课堂操作门槛'),
        ('长期学习记忆','保存会话并提取稳定学习信息','辅助连续学习与复习建议'),
        ('课堂答题','点击、语音、手势作答并获得解释','形成即时互动反馈')
    ],[3.0,6.0,7.2])

    page_break(doc)
    add_heading(doc,'2 多Agent智能体设计',1)
    add_body(doc,'数智课堂采用“总调度Agent＋专业Agent”的协同设计。小智总调度Agent是统一入口，专业Agent分别承担理解规划、演示执行、知识讲解和课堂追问职责。')
    add_table(doc,['Agent','主要职责','典型输出'],[
        ('小智总调度Agent','理解自然语言需求，判断教学演示、模型控制、课堂答题或建模任务','任务类型、简短回应、后续处理请求'),
        ('理解规划Agent','分析教学目标，匹配模型并生成结构化演示步骤','教学主题、模型、步骤与工具清单'),
        ('演示执行Agent','按照计划调用受控3D教学工具','模型加载、旋转、缩放、拆解、复位'),
        ('知识讲解Agent','结合模型和教学目标生成中学生易懂的讲解','200—400字知识讲解与语音播报'),
        ('活泼追问Agent','根据当前模型生成双选项问题并反馈','问题、A/B选项、答案与解释')
    ],[3.4,7.4,5.4])
    callout(doc,'协作原则','总调度负责“判断做什么”，规划负责“安排怎么做”，执行负责“调用工具做出来”，讲解负责“把知识讲清楚”，追问负责“检查理解并鼓励参与”。','success')
    add_figure(doc,agent_flow,'图1  数智课堂多Agent协同与长期学习记忆流程',16.0)

    page_break(doc)
    add_heading(doc,'2.1 任务路由与工具安全',2)
    add_body(doc,'系统使用规则指令处理“放大、缩小、旋转、停止”等明确控制需求，使用大模型理解复杂教学任务。教学模型名称与建模请求被严格区分，只有用户明确要求创建或生成3D模型时才进入建模工作台。')
    add_heading(doc,'受控工具范围',3)
    add_table(doc,['工具能力','作用','限制'],[
        ('加载模型','切换到指定教学模型','只能使用受支持的教学模型'),('自动旋转','从不同角度观察模型','可被停止或旋转锁定'),('自动缩放','拉近或拉远视角','按预设时间和速度执行'),('模型拆解','展示内部层次或部件关系','部分完整结构模型禁止拆解'),('复位布局','恢复模型初始布局','地球分层演示按教学规则保持分离'),('启用手势','打开摄像头手势控制','需要浏览器授权且需HTTPS或localhost')
    ],[3.0,6.2,7.0])
    callout(doc,'安全说明','学习记忆以只读背景提供给总调度Agent，不能被当作工具指令，也不能扩大Agent的操作权限。','risk')

    page_break(doc)
    add_heading(doc,'3 长期学习记忆',1)
    add_heading(doc,'3.1 工作机制',2)
    add_body(doc,'长期学习记忆按普通用户账号隔离。开启后，系统逐轮保存课堂中的用户输入、小智响应、知识讲解和追问答题事件，并在后台异步生成会话摘要，提取稳定且有教学价值的信息。')
    add_table(doc,['记忆类别','示例','用途'],[
        ('个人信息','学习阶段、称呼等稳定信息','调整交流方式'),('学习偏好','偏好图示、希望讲解更慢','改善讲解风格'),('已学主题','已学习地球内部结构','提供复习衔接'),('薄弱点','容易混淆外核与内核','提示重点复习'),('掌握情况','能够正确识别心脏四腔','辅助选择讲解深度')
    ],[3.0,6.4,6.8])
    add_heading(doc,'3.2 保存与整理规则',2)
    bullet(doc,'30分钟内持续活动会沿用最近课堂会话；超过空闲窗口后创建新会话。')
    bullet(doc,'每累计10轮用户对话，或会话空闲后，后台整理摘要和长期记忆。')
    bullet(doc,'原始会话内容保留30天；摘要和长期记忆保留至用户主动删除。')
    bullet(doc,'新发生的课堂内容可能需要等待后台整理后才出现在“学习记忆”列表。')
    callout(doc,'用户知情与控制','首次启用时系统展示说明。用户可以随时关闭长期记忆、修改记忆内容、删除单条记忆或清空全部记忆。','info')

    page_break(doc)
    add_heading(doc,'3.3 管理学习记忆',2)
    bullet(doc,'点击课堂工作台右上角的用户头像。',True)
    bullet(doc,'在账号菜单中选择“学习记忆”。',True)
    bullet(doc,'查看“长期学习记忆”开关；关闭后不再创建或追加学习会话。',True)
    bullet(doc,'在“小智记住的内容”区域查看记忆类别、内容、更新时间和来源摘要。',True)
    bullet(doc,'点击“编辑”修正内容，或点击“删除”移除单条记忆。',True)
    bullet(doc,'需要删除全部长期记忆时，点击右上方“清空全部”。',True)
    callout(doc,'注意','关闭长期记忆与删除已有记忆是两个不同操作。关闭开关停止后续记录；如需移除既有内容，还应删除单条记忆或执行“清空全部”。','warn')
    add_heading(doc,'记忆如何影响后续交互',2)
    add_body(doc,'当用户再次向小智提出问题时，服务端会选取账号下置信度较高、更新时间较新的相关记忆，并与当前会话摘要一起作为只读学习背景交给小智总调度Agent。总调度Agent可据此调整讲解难度或提出复习建议，其他工具权限保持不变。')

    page_break(doc)
    add_heading(doc,'4 访问、注册与登录',1)
    add_heading(doc,'4.1 访问环境',2)
    add_table(doc,['项目','要求'],[
        ('访问地址','【请填写正式访问地址】'),('浏览器','最新版 Chrome 或 Microsoft Edge'),('网络','可访问系统服务及配置的大模型服务'),('摄像头','手势功能需要摄像头；远程部署需HTTPS'),('麦克风','语音输入需要浏览器授权'),('显示设备','建议1920×1080或更高分辨率')
    ],[4.0,12.2])
    add_heading(doc,'4.2 普通用户登录',2)
    bullet(doc,'打开访问地址，点击进入课堂。',True); bullet(doc,'选择“用户登录”，输入用户名和密码。',True); bullet(doc,'首次使用可点击“注册”，创建普通用户账号后直接进入课堂。',True)
    add_body(doc,'用户名支持中文、字母、数字、下划线、短横线或邮箱地址；密码长度为6—128个字符。')
    add_heading(doc,'4.3 管理员登录',2)
    add_body(doc,'在登录页切换到“管理员”，使用由系统管理员创建的管理员账号登录。管理员登录后进入用户管理后台，不进入普通用户的学习记忆流程。')

    page_break(doc)
    add_heading(doc,'5 课堂工作台',1)
    add_figure(doc,workspace,'图2  课堂工作台功能分区示意',16.0)
    add_table(doc,['区域','主要功能'],[
        ('顶部导航','3D建模生成、导入模型、个人设置、学习记忆、退出登录'),('学科资源库','浏览“我的模型”、建模资源及化学、生物、地理模型'),('3D模型展示区','模型观察、知识讲解、全屏、答题模式和结构图查看'),('小智助教','显示语音状态、回应和当前教学提示'),('多智能体协作台','输入教学需求，查看Agent状态、步骤和工具执行结果'),('底部控制区','语音开关、交互速度设置和答题入口')
    ],[4.2,12.0])

    page_break(doc)
    add_heading(doc,'6 模型选择与基本操作',1)
    add_heading(doc,'6.1 选择内置模型',2)
    bullet(doc,'展开左侧“学科资源库”。',True); bullet(doc,'选择化学、生物或地理分类。',True); bullet(doc,'点击模型名称，等待加载提示结束。',True); bullet(doc,'模型出现后使用鼠标拖动旋转，滚轮或触控手势缩放。',True)
    add_heading(doc,'6.2 导入自有模型',2)
    bullet(doc,'点击顶部“导入模型”。',True); bullet(doc,'选择GLB、GLTF或FBX模型文件及必要的纹理文件。',True); bullet(doc,'等待系统完成读取，确认文件名显示在模型区。',True)
    callout(doc,'兼容性提示','不同模型的内部层级不同，自动拆解效果取决于模型结构。金刚石等完整结构展示模型不支持拆解。','warn')
    add_heading(doc,'6.3 全屏与交互速度',2)
    add_body(doc,'点击模型区右上角全屏按钮进入展示区全屏；点击底部设置按钮可分别调整缩放速度和旋转速度，也可以使用“重置默认”恢复系统预设。')

    page_break(doc)
    add_heading(doc,'7 语音助教与多Agent演示',1)
    add_heading(doc,'7.1 使用小智语音助教',2)
    bullet(doc,'先加载一个支持交互的3D模型。',True); bullet(doc,'点击底部麦克风按钮并允许浏览器使用麦克风。',True); bullet(doc,'清楚说出需求，例如“讲解地球内部结构”或“把模型放大一点”。',True); bullet(doc,'等待小智回应；再次点击麦克风可关闭语音输入。',True)
    add_table(doc,['任务类型','示例指令','系统行为'],[
        ('教学演示','讲解地球内部结构','规划步骤、执行模型操作并生成讲解'),('模型控制','放大一点／转一转','快速调用缩放或旋转工具'),('课堂追问','考考我／开始答题','直接启动活泼追问Agent'),('3D建模','生成一个苹果3D模型','进入3D建模工作台')
    ],[3.2,5.0,8.0])
    add_heading(doc,'7.2 多智能体协作台',2)
    bullet(doc,'切换左侧栏到“多智能体平台”。',True); bullet(doc,'在输入框填写教学需求，按Enter启动，Shift+Enter换行。',True); bullet(doc,'观察五类Agent的待命、规划中、执行中、完成或异常状态。',True); bullet(doc,'查看时间线中的演示步骤、讲解摘要和工具执行结果。',True)

    page_break(doc)
    add_heading(doc,'8 手势交互',1)
    add_heading(doc,'8.1 开启手势捕捉',2)
    bullet(doc,'加载本地GLB教学模型。URL交互页面不支持本地手势捕捉。',True); bullet(doc,'点击左侧手势按钮，允许浏览器使用摄像头。',True); bullet(doc,'等待摄像头预览出现，并让手掌位于画面中央。',True)
    add_heading(doc,'8.2 单手模式',2)
    add_table(doc,['手势','作用','操作要点'],[
        ('张开右手','放大模型','保持手掌完整入镜'),('右手握拳','缩小模型','动作清晰并短暂保持'),('食指＋拇指捏合','拖拽模型零件','捏合后移动手掌'),('食指＋中指并拢滑动','旋转画面','双指旋转优先于缩放')
    ],[4.0,4.2,8.0])
    add_heading(doc,'8.3 双手模式',2)
    add_body(doc,'双手模式下，左手张开或闭合负责缩放，右手负责旋转和零件拖拽。缩放与拖拽互斥，避免同一动作触发多个控制。')
    callout(doc,'识别建议','保持光线均匀、摄像头稳定、手掌与背景有明显对比；手势没有反应时先检查浏览器摄像头权限。','info')

    page_break(doc)
    add_heading(doc,'9 课堂答题与活泼追问Agent',1)
    add_heading(doc,'9.1 启动方式',2)
    bullet(doc,'知识讲解完成后，系统可自动生成一道轻量追问题。')
    bullet(doc,'用户可直接说“考考我”“开始答题”或“出一道题”。')
    bullet(doc,'也可点击模型区的“答题模式”，进入五题课堂测验。')
    add_heading(doc,'9.2 作答方式',2)
    add_table(doc,['方式','操作','确认规则'],[
        ('点击','直接点击A或B选项','立即提交'),('语音','说“A”“B”或对应选项内容','识别最终语音后提交'),('手势','移动手掌指针到选项区域','稳定悬停1.2秒确认')
    ],[3.2,7.0,6.0])
    add_heading(doc,'9.3 反馈与记忆',2)
    add_body(doc,'系统完成判题后展示正确答案、知识解释和鼓励性反馈，约3秒后返回模型界面。若长期学习记忆已开启，本轮答题结果会先作为课堂事件保存，并在后续异步整理时可能形成“薄弱点”或“掌握情况”记忆。')
    callout(doc,'准确表述','不是每次答题都会立刻生成长期记忆；系统只提取稳定且具有教学价值的信息。','warn')

    page_break(doc)
    add_heading(doc,'10 3D建模与“我的模型”',1)
    bullet(doc,'点击顶部“3D建模生成”。',True); bullet(doc,'输入模型描述，或上传用于图生3D的参考图片。',True); bullet(doc,'点击“开始生成3D模型”，生成过程通常需要一定时间。',True); bullet(doc,'生成完成后预览结果，点击“一键导入”。',True); bullet(doc,'返回课堂，在“我的模型”中打开已保存模型。',True)
    callout(doc,'课堂建议','生成任务可在后台继续运行。教师可以先返回课堂使用现有模型，待生成完成后再导入。','info')
    add_heading(doc,'模型生成与内置模型的区别',2)
    add_body(doc,'“打开或讲解心脏模型”属于内置教学模型任务；只有“创建、生成、建模”等明确表达才会进入3D建模工作台，避免误操作。')

    page_break(doc)
    add_heading(doc,'11 典型案例：地球内部结构',1)
    add_figure(doc,seq,'图3  “讲解地球内部结构”多Agent任务时序',16.0)
    bullet(doc,'教师说：“讲解地球内部结构，展示地壳、地幔、外核和内核的关系。”',True)
    bullet(doc,'小智总调度Agent识别为教学演示任务，并交给理解规划Agent。',True)
    bullet(doc,'理解规划Agent匹配地球内部结构模型，生成加载、观察、拆解和讲解步骤。',True)
    bullet(doc,'演示执行Agent调用模型工具，保持四层分离以便观察。',True)
    bullet(doc,'知识讲解Agent根据当前模型生成并播报知识内容。',True)
    bullet(doc,'活泼追问Agent可生成A/B题，学生通过点击、语音或手势作答。',True)
    bullet(doc,'长期记忆开启时，会话与答题事件被保存；后台可能提取“已学习地球内部结构”或相关薄弱点。',True)

    page_break(doc)
    add_heading(doc,'12 教学价值示范',1)
    add_table(doc,['价值维度','传统方式痛点','数智课堂体现'],[
        ('空间认知','二维图片难以呈现结构关系','3D模型支持旋转、缩放和拆解'),('课堂操作','教师需频繁切换鼠标键盘','语音、手势和触控协同'),('任务组织','教学演示步骤依赖人工准备','规划Agent生成结构化演示路线'),('即时讲解','固定素材难以响应临时问题','知识讲解Agent按模型和目标生成内容'),('互动反馈','课堂提问反馈不及时','追问Agent提供即时判题与解释'),('连续学习','不同课堂之间缺乏衔接','长期记忆辅助后续难度与复习建议')
    ],[3.0,6.1,7.1])
    add_body(doc,'系统的核心价值不只是“展示3D模型”，而是通过多Agent专业分工，把需求理解、自主规划、工具执行、知识讲解、互动反馈和长期学习记忆连接为可观察、可控制的课堂流程。')

    page_break(doc)
    add_heading(doc,'13 常见问题',1)
    add_table(doc,['问题','处理方法'],[
        ('无法登录','检查账号密码、账号是否被禁用以及后端服务是否可用。'),('模型加载失败','刷新页面，检查网络和模型文件格式；大型模型需要更长加载时间。'),('语音不可用','使用Chrome或Edge，并在地址栏权限中允许麦克风。'),('摄像头无法开启','确认使用HTTPS或localhost，并允许摄像头权限。'),('手势识别不稳定','改善光线和背景对比，让完整手掌进入画面。'),('Agent提示异常','检查网络与大模型服务配置；本地模型仍可手动演示。'),('记忆没有立即出现','长期记忆异步整理，请完成更多对话或等待会话空闲。'),('不希望保存学习记录','进入“学习记忆”关闭开关，并按需删除已有记忆。')
    ],[4.5,11.7])

    page_break(doc)
    add_heading(doc,'附录A 个人设置与退出登录',1)
    add_body(doc,'点击右上角用户头像可进入个人设置、学习记忆或退出登录。个人设置支持修改昵称、上传PNG/JPEG/WebP头像和移除头像。头像在保存前会自动压缩。')
    add_heading(doc,'附录B 管理员用户管理',1)
    bullet(doc,'管理员登录后进入“用户管理后台”。')
    bullet(doc,'可刷新用户列表，查看用户名、角色、状态和创建时间。')
    bullet(doc,'可将普通用户升为管理员，或将管理员降为普通用户。')
    bullet(doc,'可启用或禁用其他账号；不能修改自身角色或禁用当前账号。')
    add_heading(doc,'附录C 隐私与数据说明',1)
    bullet(doc,'学习记忆仅面向普通用户账号，并按账号隔离。')
    bullet(doc,'原始会话保留30天；摘要与长期记忆保留至用户删除。')
    bullet(doc,'用户可以查看、编辑、删除、清空或关闭长期记忆。')
    bullet(doc,'记忆是辅助教学的学习背景，不是正式成绩或学情评价。')
    callout(doc,'配置提醒','提交申报材料前，请将封面和文档信息页中的访问地址、国内部署平台占位符替换为正式信息。','warn')

    # Update fields on open
    settings=doc.settings._element; update=OxmlElement('w:updateFields'); update.set(qn('w:val'),'true'); settings.append(update)
    doc.core_properties.title='数智课堂教育智能体操作手册 V1.0'; doc.core_properties.subject='数智课堂多Agent与长期学习记忆操作说明'; doc.core_properties.author='数智课堂项目组'
    doc.save(DOCX)
    print(DOCX)


if __name__=='__main__': build()
